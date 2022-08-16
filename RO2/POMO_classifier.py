import numpy as np
import pandas as pd
pd.pandas.set_option('display.max_columns',None)

import seaborn as sns
sns.set(font_scale=1.2)

import matplotlib.pyplot as plt
plt.rcParams['figure.figsize'] = (12,8)
# %matplotlib inline

from sklearn.model_selection import train_test_split

from sklearn import metrics
from sklearn.metrics import accuracy_score
# from sklearn.metrics import log_loss
# from sklearn.metrics import cohen_kappa_score
# from sklearn.metrics import precision_score
# from sklearn.metrics import recall_score
# from sklearn.metrics import f1_score
from sklearn.metrics import zero_one_loss
from sklearn.metrics import hamming_loss


from sklearn.metrics import multilabel_confusion_matrix # confusion_matrix

import pickle
import re
import logging

# from sklearn.linear_model import LogisticRegression  
# from sklearn.neighbors import KNeighborsClassifier 
# from sklearn.naive_bayes import GaussianNB
# from sklearn.svm import SVC
# from sklearn.tree import DecisionTreeClassifier
# from sklearn.ensemble import RandomForestClassifier
# from sklearn.ensemble import GradientBoostingClassifier
# from sklearn.ensemble import AdaBoostClassifier
from sklearn.multioutput import MultiOutputClassifier

import docx
from docx.shared import Inches
import datetime
import os

import warnings
warnings.filterwarnings("ignore", category=FutureWarning)


ROOT_PATH = os.getcwd()
ARTIFACT_DIR = "Artifact"
ROOT_DIR = os.path.join(ROOT_PATH, ARTIFACT_DIR)
CURRENT_TIME_STAMP = f"{datetime.datetime.now().strftime('%Y-%m-%d-%H-%M-%S')}"

MODELS_DIR = "Models"
IMAGES_DIR = "Images"
DOCS_DIR = "Docs"
LOG_DIR = "Model_Logs"
LOG_FILE_NAME = f"Log_{CURRENT_TIME_STAMP}.log"
RESULT_FILE_NAME = "MultiLabelModel_result.csv"
DOCS_FILE_NAME = "Classification_Report.docx"

MODELS_DIR_PATH = os.path.join(ROOT_DIR, CURRENT_TIME_STAMP, MODELS_DIR)
IMAGES_DIR_PATH = os.path.join(ROOT_DIR, CURRENT_TIME_STAMP, IMAGES_DIR)
DOCS_DIR_PATH = os.path.join(ROOT_DIR, CURRENT_TIME_STAMP, DOCS_DIR)
LOG_DIR_PATH = os.path.join(ROOT_DIR, CURRENT_TIME_STAMP, LOG_DIR)
LOG_FILE_PATH = os.path.join(LOG_DIR_PATH, LOG_FILE_NAME)
RESULT_FILE_PATH = os.path.join(DOCS_DIR_PATH, RESULT_FILE_NAME)
DOCS_FILE_PATH = os.path.join(DOCS_DIR_PATH, DOCS_FILE_NAME)

os.makedirs(MODELS_DIR_PATH, exist_ok = True)
os.makedirs(IMAGES_DIR_PATH, exist_ok = True)
os.makedirs(DOCS_DIR_PATH, exist_ok = True)
os.makedirs(LOG_DIR_PATH, exist_ok = True)

logging.basicConfig(
    filename=LOG_FILE_PATH, 
    filemode="w", 
    format='[%(asctime)s] %(name)s - %(levelname)s - %(message)s', 
    level=logging.INFO
    )

doc = docx.Document()
doc.add_heading(f'Classification Report {CURRENT_TIME_STAMP}', 0)


def builMultiLabelModels(train, test, classifierModel):
    
    logging.info(f"Initializing builMultiLabelModels")
    doc.add_paragraph(f"Initializing builMultiLabelModels")
    
    X_train, y_train = train
    X_test, y_test = test

    multiLabelModel = MultiOutputClassifier(classifierModel, n_jobs=2)
    
    classifierModel = multiLabelModel.fit(X_train, y_train)
    
    y_pred = classifierModel.predict(X_train)
    
    Accuracy_score = round((accuracy_score(y_train, y_pred))*100,2)
    Loss = round((1-accuracy_score(y_train, y_pred))*100,2)
    Zero_one_loss = round(zero_one_loss(y_train, y_pred)*100,2)
    Hamming_loss = round((hamming_loss(y_train, y_pred))*100,2)
    
    doc.add_paragraph(f"Classification on Training dataset : ")
    doc.add_paragraph(f"Accuracy_score: {Accuracy_score} %")
    doc.add_paragraph(f"Loss: {Loss} %")
    doc.add_paragraph(f"Zero_one_loss: {Zero_one_loss} %")
    doc.add_paragraph(f"Hamming_loss: {Hamming_loss} %")
    doc.add_paragraph(f"Classification_report: {metrics.classification_report(y_train, y_pred)}")

    logging.info(f"Classification on Training dataset : [ 'Accuracy': {Accuracy_score}, 'Loss': {Loss}, 'Zero_one_loss': {Zero_one_loss}, 'Hamming_loss': {Hamming_loss} ]")


    classifierModel = multiLabelModel.fit(X_test, y_test)
    
    y_pred = classifierModel.predict(X_test)
    
    Accuracy_score_ = round((accuracy_score(y_test, y_pred))*100,2)
    Loss_ = round((1-accuracy_score(y_test, y_pred))*100,2)
    Zero_one_loss_ = round(zero_one_loss(y_test, y_pred)*100,2)
    Hamming_loss_ = round((hamming_loss(y_test, y_pred))*100,2)
    
    doc.add_paragraph(f"Classification on Testing dataset : ")
    doc.add_paragraph(f"Accuracy_score: {Accuracy_score_} %")
    doc.add_paragraph(f"Loss: {Loss_} %")
    doc.add_paragraph(f"Zero_one_loss: {Zero_one_loss_} %")
    doc.add_paragraph(f"Hamming_loss: {Hamming_loss_} %")
    doc.add_paragraph(f"Classification_report: {metrics.classification_report(y_test, y_pred)}")

    logging.info(f"Classification on Testing dataset : [ 'Accuracy': {Accuracy_score_}, 'Loss': {Loss_}, 'Zero_one_loss': {Zero_one_loss_}, 'Hamming_loss': {Hamming_loss_} ]")


    logging.info(f"Done builMultiLabelModels")
    return classifierModel, {'Accuracy': Accuracy_score, 'Loss': Loss, 'Hamming_loss': Hamming_loss}

def predict_ensemble_voting_classifier(X_test, path : str = RESULT_FILE_PATH):
    
    # selecting top 3 models
    if os.path.exists(path):   

        models_df = pd.read_csv(path)
        models_list = list(models_df.sort_values(by = ['Accuracy'], ascending = False)[0:3]['Filename'].values)

        # get model prediction results
        output_list = []
        for path in models_list:
            model = pickle.load(open(path, 'rb')) 
            pred = model.predict(X_test)
            output_list.append(pred)
        
        # get predicted values (from votting)
        and_of_pred = (output_list[0] & output_list[1] & output_list[2])
        or_of_pred = (output_list[0] | output_list[1] | output_list[2])
        y_pred_new = np.nan_to_num(and_of_pred // or_of_pred)
        
        return y_pred_new

    else:
        raise Exception(f"[ {RESULT_FILE_PATH} ] : file not found.")

def classification_report_ensemble_voting_classifier(y_train, y_pred_new):

    Accuracy_score = round((accuracy_score(y_train, y_pred_new))*100,2)
    Loss = round((1-accuracy_score(y_train, y_pred_new))*100,2)
    Zero_one_loss = round(zero_one_loss(y_train, y_pred_new)*100,2)
    Hamming_loss = round((hamming_loss(y_train, y_pred_new))*100,2)
    # ML_confusion_matrix = multilabel_confusion_matrix(y_train, y_pred_new)
    ML_confusion_matrix = metrics.classification_report(y_train, y_pred_new)
    return (Accuracy_score, Loss, Zero_one_loss, Hamming_loss, ML_confusion_matrix)


def fit_ensemble_voting_classifier(train, test, classifiers):
    
    logging.info("========== Initializing fit_ensemble_voting_classifier ==========\n\n")
    doc.add_heading(f"========== Initializing fit_ensemble_voting_classifier ==========\n\n", 2)
    result_df = pd.DataFrame()

    for classifier in classifiers:

        # X_train, y_train, X_test, y_test  = train, test

        logging.info(f"Model : {type(classifier).__name__} running...")
        # print(f'Model : {type(classifier).__name__}')
        doc.add_heading(f"Model : {type(classifier).__name__}\n", 2)

        model, data = builMultiLabelModels(train, test, classifier)

        # save the model to disk
        FILE_NAME = f'MultiLabelModel_{type(classifier).__name__}'
        FILE_NAME = re.sub('\W+','_', FILE_NAME )+'.pkl'
        MODELS_FILE_PATH = os.path.join(MODELS_DIR_PATH, FILE_NAME)
        pickle.dump(model, open(MODELS_FILE_PATH, 'wb'))

        # print(f'File saved : {MODELS_FILE_PATH}')    
        logging.info(f"Model file saved : {MODELS_FILE_PATH}")
        doc.add_paragraph(f"File saved: {MODELS_FILE_PATH}")

        data['Model_name'] = type(classifier).__name__
        data['Filename'] = MODELS_FILE_PATH

        # print(f'data : {data}')
        logging.info(f"Model data : {data} \n\n")
        doc.add_paragraph(f"Model data: {data} \n\n")

        result_df = result_df.append(data, ignore_index=True)

    result_df1 = result_df[['Model_name', 'Accuracy', 'Loss', 'Hamming_loss', 'Filename']]
    result_df1.to_csv(RESULT_FILE_PATH, index = False)

    logging.info(f"Result file saved : {RESULT_FILE_PATH}")

    try:
        y_pred_new = predict_ensemble_voting_classifier(train[0], RESULT_FILE_PATH)

        Accuracy_score, Loss, Zero_one_loss, Hamming_loss, ML_confusion_matrix = classification_report_ensemble_voting_classifier(train[1], y_pred_new)
        
        print("Classification Report on Training dataset : ")
        print("Accuracy_score:", Accuracy_score,'%')
        print("Loss:", Loss,'%')
        print("Zero_one_loss:", Zero_one_loss,'%')
        print("Hamming_loss:", Hamming_loss,'%')
        print("ML_confusion_matrix:\n", ML_confusion_matrix)

        doc.add_heading(f"Ensemble_Voting_Classifier_Result >>> \n", 2)
        doc.add_paragraph(f"Classification Report on Training dataset : ")
        doc.add_paragraph(f"Accuracy_score: {Accuracy_score} %")
        doc.add_paragraph(f"Loss: {Loss} %")
        doc.add_paragraph(f"Zero_one_loss: {Zero_one_loss} %")
        doc.add_paragraph(f"Hamming_loss: {Hamming_loss} %")
        doc.add_paragraph(f"ML_confusion_matrix:\n{ML_confusion_matrix}")


        y_pred_new = predict_ensemble_voting_classifier(test[0], RESULT_FILE_PATH)

        Accuracy_score, Loss, Zero_one_loss, Hamming_loss, ML_confusion_matrix = classification_report_ensemble_voting_classifier(test[1], y_pred_new)
        
        print("Classification Report on Testing dataset : ")
        print("Accuracy_score:", Accuracy_score,'%')
        print("Loss:", Loss,'%')
        print("Zero_one_loss:", Zero_one_loss,'%')
        print("Hamming_loss:", Hamming_loss,'%')
        print("ML_confusion_matrix:\n", ML_confusion_matrix)

        doc.add_paragraph(f"Classification Report on Testing dataset : ")
        doc.add_paragraph(f"Accuracy_score: {Accuracy_score} %")
        doc.add_paragraph(f"Loss: {Loss} %")
        doc.add_paragraph(f"Zero_one_loss: {Zero_one_loss} %")
        doc.add_paragraph(f"Hamming_loss: {Hamming_loss} %")
        doc.add_paragraph(f"ML_confusion_matrix:\n{ML_confusion_matrix}")

    except Exception as e:
        logging.error(f"Exception : [ {e} ]")
        print(f"Exception : [ {e} ]")
    

    doc.save(DOCS_FILE_PATH)    
    logging.info("========== Done fit_ensemble_voting_classifier ==========\n\n")
    return {'RESULT_FILE_PATH': RESULT_FILE_PATH}

